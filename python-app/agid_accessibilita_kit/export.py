from __future__ import annotations

import pathlib


def export_docx(cfg: dict, out_path: str) -> pathlib.Path:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Dipendenza mancante: python-docx. Esegui: python -m pip install -e .") from exc

    ente = cfg.get("ente", "Nome organizzazione")
    sito = cfg.get("sito", "https://www.esempio.it")
    email = cfg.get("email", "accessibilita@esempio.it")
    base_path = cfg.get("basePath", "/accessibilita")

    doc = Document()
    doc.add_heading("Accessibilita (bozze)", level=1)
    doc.add_paragraph(f"Organizzazione: {ente}")
    doc.add_paragraph(f"Sito/app: {sito}")
    doc.add_paragraph("")

    doc.add_heading("Dichiarazione di accessibilita (bozza)", level=2)
    doc.add_paragraph(
        "Questa e' una bozza. Compila e pubblica la dichiarazione secondo le procedure ufficiali applicabili."
    )
    doc.add_paragraph(f"Contatto accessibilita: {email}")
    doc.add_paragraph(f"Percorso pubblico consigliato: {base_path}")
    doc.add_paragraph("")

    doc.add_heading("Feedback accessibilita (bozza)", level=2)
    doc.add_paragraph("Per segnalazioni di accessibilita, contattaci indicando:")
    for item in [
        "URL della pagina",
        "Dispositivo e browser",
        "Tecnologia assistiva (se usata)",
        "Descrizione del problema",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph("")

    doc.add_heading("Snippet footer (bozza)", level=2)
    doc.add_paragraph("Inserisci nel footer del sito un link chiaro a:")
    doc.add_paragraph(f"{base_path}/dichiarazione-accessibilita.html", style="List Bullet")
    doc.add_paragraph(f"{base_path}/feedback-accessibilita.html", style="List Bullet")
    doc.add_paragraph("")

    doc.add_heading("Checklist operativa (bozza)", level=2)
    for item in [
        "Link nel footer a Dichiarazione di accessibilita",
        "Canale feedback (email/form) funzionante",
        "Attributo lang su <html>",
        "Presenza di un <h1> principale",
        "Immagini con alt appropriato",
        "Form con label/aria-label",
        "Link e pulsanti con nome accessibile",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    p = pathlib.Path(out_path).resolve()
    p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(p))
    return p


def export_html_from_docx(docx_path: str, out_path: str) -> pathlib.Path:
    try:
        import mammoth  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Dipendenza mancante: mammoth. Esegui: python -m pip install -e .") from exc

    docx_p = pathlib.Path(docx_path).resolve()
    out_p = pathlib.Path(out_path).resolve()
    out_p.parent.mkdir(parents=True, exist_ok=True)

    with open(docx_p, "rb") as f:
        result = mammoth.convert_to_html(f)
    html = result.value

    out_p.write_text(
        "<!doctype html><html lang=\"it\"><head><meta charset=\"utf-8\" />"
        "<meta name=\"viewport\" content=\"width=device-width, initial-scale=1\" />"
        "<title>Accessibilita (bozze)</title></head><body>"
        + html
        + "</body></html>\n",
        encoding="utf-8",
    )
    return out_p


def export_pdf_from_docx(docx_path: str, out_path: str) -> pathlib.Path:
    # docx2pdf usa Microsoft Word su Windows: se Word non e' installato, fallira'.
    try:
        from docx2pdf import convert  # type: ignore
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("Dipendenza mancante: docx2pdf. Esegui: python -m pip install -e .") from exc

    docx_p = pathlib.Path(docx_path).resolve()
    out_p = pathlib.Path(out_path).resolve()
    out_p.parent.mkdir(parents=True, exist_ok=True)

    try:
        convert(str(docx_p), str(out_p))
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(
            "Export PDF fallito. Su Windows, docx2pdf in genere richiede Microsoft Word installato."
        ) from exc
    return out_p
